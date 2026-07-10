"""解析 .docx 文本、段落、表格。

按文档实际顺序（段落与表格交错排列）提取所有内容块。
"""

from dataclasses import dataclass, field
from pathlib import Path


@dataclass
class ParsedBlock:
    """单个解析块。"""

    block_id: str
    doc_id: str
    block_index: int
    block_type: str  # "paragraph" | "table_row"
    text: str
    paragraph_index: int | None = None
    table_index: int | None = None
    row_index: int | None = None
    cell_index: int | None = None
    style_name: str | None = None
    detected_level: str | None = None
    article_no: str | None = None


@dataclass
class ParsedDocument:
    """解析后的文档。"""

    doc_id: str
    blocks: list[ParsedBlock] = field(default_factory=list)
    parse_error: str | None = None
    paragraph_count: int = 0
    table_count: int = 0


def parse_docx(path: Path, doc_id: str) -> ParsedDocument:
    """解析一个 .docx 文件，按文档顺序返回所有结构化块。

    段落和表格按照在文档 body 中的实际出现顺序交错排列。

    Args:
        path: .docx 文件路径
        doc_id: 文档唯一 ID

    Returns:
        ParsedDocument，包含所有块的列表；解析失败时 parse_error 非空。
    """
    try:
        from docx import Document
    except ImportError:
        return ParsedDocument(
            doc_id=doc_id,
            parse_error="python-docx not installed",
        )

    result = ParsedDocument(doc_id=doc_id)

    try:
        doc = Document(str(path))
    except Exception as e:
        result.parse_error = f"Failed to open document: {e}"
        return result

    blocks: list[ParsedBlock] = []

    try:
        body = doc.element.body
        para_idx = 0
        table_idx = 0
        paragraph_count = 0

        for child in body:
            # 提取本地标签名（去掉命名空间前缀）
            tag = _local_tag(child.tag)

            if tag == "p":
                # ── 段落 ──
                if para_idx < len(doc.paragraphs):
                    para = doc.paragraphs[para_idx]
                    text = para.text
                    if text and text.strip():
                        blocks.append(ParsedBlock(
                            block_id="",
                            doc_id=doc_id,
                            block_index=0,  # 后续重新编号
                            block_type="paragraph",
                            text=text.strip(),
                            paragraph_index=para_idx,
                            style_name=para.style.name if para.style else None,
                        ))
                        paragraph_count += 1
                para_idx += 1

            elif tag == "tbl":
                # ── 表格 ──
                if table_idx < len(doc.tables):
                    table = doc.tables[table_idx]
                    _parse_single_table(blocks, table, doc_id, table_idx)
                table_idx += 1

            # 其他元素（图片、分页符等）跳过

        # 重新编号
        for i, b in enumerate(blocks):
            b.block_index = i

        result.blocks = blocks
        result.paragraph_count = paragraph_count
        result.table_count = table_idx

    except Exception as e:
        result.parse_error = f"Parse error: {e}"

    return result


def parse_paragraphs(document) -> list[ParsedBlock]:
    """解析普通段落（辅助函数，供外部单独调用）。"""
    blocks: list[ParsedBlock] = []
    for i, para in enumerate(document.paragraphs):
        text = para.text
        if not text or not text.strip():
            continue
        block = ParsedBlock(
            block_id="",
            doc_id="",
            block_index=i,
            block_type="paragraph",
            text=text.strip(),
            paragraph_index=i,
            style_name=para.style.name if para.style else None,
        )
        blocks.append(block)
    return blocks


def parse_tables(document) -> list[ParsedBlock]:
    """解析表格内容（辅助函数，供外部单独调用）。"""
    blocks: list[ParsedBlock] = []
    for t_idx, table in enumerate(document.tables):
        _parse_single_table(blocks, table, "", t_idx)
    return blocks


def _parse_single_table(
    blocks: list[ParsedBlock], table, doc_id: str, table_idx: int,
) -> None:
    """解析单个表格，生成合并行文本块。"""
    for r_idx, row in enumerate(table.rows):
        row_parts: list[str] = []
        for c_idx, cell in enumerate(row.cells):
            cell_text = cell.text.strip()
            if cell_text:
                row_parts.append(f"列{c_idx + 1}：{cell_text}")

        if row_parts:
            block = ParsedBlock(
                block_id="",
                doc_id=doc_id,
                block_index=0,
                block_type="table_row",
                text="表格行：" + " | ".join(row_parts),
                table_index=table_idx,
                row_index=r_idx,
            )
            blocks.append(block)


def _local_tag(tag: str) -> str:
    """从带命名空间的 XML 标签中提取本地名称。

    '{http://...}p' → 'p'
    'p' → 'p'
    """
    if "}" in tag:
        return tag.split("}", 1)[1]
    return tag
